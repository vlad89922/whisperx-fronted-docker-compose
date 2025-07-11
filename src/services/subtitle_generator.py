"""
Сервис для генерации файлов субтитров в различных форматах
"""
from typing import List, Dict, Optional
from pathlib import Path
from datetime import datetime

from ..config.settings import TRANSCRIPTS_DIR, TEMP_DIR
from ..utils.time_formatters import format_time_srt, format_time_vtt, format_time_tsv

# Проверяем доступность библиотек
try:
    from docx import Document
    from docx.shared import Inches, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class SubtitleGenerator:
    """Генератор файлов субтитров"""
    
    @staticmethod
    def generate_srt(segments: List[Dict], task_id: str, filename: str, temp: bool = False) -> Optional[str]:
        """Генерация SRT файла"""
        content = []
        for i, segment in enumerate(segments, 1):
            start_time = format_time_srt(segment['start'])
            end_time = format_time_srt(segment['end'])
            text = segment['text'].strip()
            
            content.append(f"{i}")
            content.append(f"{start_time} --> {end_time}")
            content.append(text)
            content.append("")  # Пустая строка между субтитрами
        
        file_dir = TEMP_DIR if temp else TRANSCRIPTS_DIR
        file_suffix = "_temp" if temp else f"_{Path(filename).stem}"
        file_path = file_dir / f"{task_id}{file_suffix}.srt"
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))
        
        return str(file_path)
    
    @staticmethod
    def generate_vtt(segments: List[Dict], task_id: str, filename: str, temp: bool = False) -> Optional[str]:
        """Генерация VTT файла"""
        content = ["WEBVTT", ""]
        for segment in segments:
            start_time = format_time_vtt(segment['start'])
            end_time = format_time_vtt(segment['end'])
            text = segment['text'].strip()
            
            content.append(f"{start_time} --> {end_time}")
            content.append(text)
            content.append("")
        
        file_dir = TEMP_DIR if temp else TRANSCRIPTS_DIR
        file_suffix = "_temp" if temp else f"_{Path(filename).stem}"
        file_path = file_dir / f"{task_id}{file_suffix}.vtt"
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))
        
        return str(file_path)
    
    @staticmethod
    def generate_tsv(segments: List[Dict], task_id: str, filename: str, temp: bool = False) -> Optional[str]:
        """Генерация TSV файла"""
        content = ["start\tend\ttext"]
        for segment in segments:
            start_time = format_time_tsv(segment['start'])
            end_time = format_time_tsv(segment['end'])
            text = segment['text'].strip().replace('\t', ' ')  # Убираем табы из текста
            
            content.append(f"{start_time}\t{end_time}\t{text}")
        
        file_dir = TEMP_DIR if temp else TRANSCRIPTS_DIR
        file_suffix = "_temp" if temp else f"_{Path(filename).stem}"
        file_path = file_dir / f"{task_id}{file_suffix}.tsv"
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))
        
        return str(file_path)
    
    @staticmethod
    def generate_docx(segments: List[Dict], task_id: str, filename: str, temp: bool = False) -> Optional[str]:
        """Генерация DOCX файла"""
        if not DOCX_AVAILABLE:
            print("❌ python-docx не доступен для создания DOCX файлов")
            return None
        
        try:
            doc = Document()
            
            # Заголовок документа
            title = doc.add_heading(f'Транскрипция: {filename}', 0)
            
            # Информация о файле
            info_para = doc.add_paragraph()
            info_para.add_run('Файл: ').bold = True
            info_para.add_run(filename)
            info_para.add_run('\nДата создания: ').bold = True
            info_para.add_run(datetime.now().strftime('%d.%m.%Y %H:%M:%S'))
            info_para.add_run('\nID задачи: ').bold = True
            info_para.add_run(task_id)
            
            doc.add_paragraph()  # Пустая строка
            
            # Добавляем сегменты
            for segment in segments:
                # Временная метка
                time_para = doc.add_paragraph()
                time_run = time_para.add_run(f"[{format_time_srt(segment['start'])} - {format_time_srt(segment['end'])}]")
                time_run.font.size = Inches(0.1)
                time_run.font.color.rgb = RGBColor(128, 128, 128)
                
                # Спикер (если есть)
                if 'speaker' in segment and segment['speaker'] is not None:
                    speaker_para = doc.add_paragraph()
                    speaker_run = speaker_para.add_run(f"Спикер {segment['speaker']}:")
                    speaker_run.bold = True
                
                # Текст
                text_para = doc.add_paragraph(segment['text'])
                text_para.add_run('\n')
            
            # Сохраняем файл
            file_dir = TEMP_DIR if temp else TRANSCRIPTS_DIR
            file_suffix = "_temp" if temp else f"_{Path(filename).stem}"
            file_path = file_dir / f"{task_id}{file_suffix}.docx"
            
            doc.save(str(file_path))
            return str(file_path)
            
        except Exception as e:
            print(f"❌ Ошибка создания DOCX: {e}")
            return None
    
    @staticmethod
    def generate_pdf(segments: List[Dict], task_id: str, filename: str, temp: bool = False) -> Optional[str]:
        """Генерация PDF файла с поддержкой UTF-8 кириллицы"""
        if not PDF_AVAILABLE:
            print("❌ reportlab не доступен для создания PDF файлов")
            return None
        
        try:
            import os
            
            file_dir = TEMP_DIR if temp else TRANSCRIPTS_DIR
            file_suffix = "_temp" if temp else f"_{Path(filename).stem}"
            file_path = file_dir / f"{task_id}{file_suffix}.pdf"
            
            # Регистрируем шрифт с поддержкой кириллицы
            font_name = None
            try:
                # Для macOS - используем системные шрифты с поддержкой кириллицы
                mac_fonts = [
                    '/Library/Fonts/Arial.ttf',
                    '/System/Library/Fonts/Helvetica.ttc',
                    '/System/Library/Fonts/Times.ttc',
                    '/Library/Fonts/Microsoft/Arial.ttf',
                    '/System/Library/Fonts/Arial Unicode MS.ttf'
                ]
                
                # Для Linux
                linux_fonts = [
                    '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
                    '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
                    '/usr/share/fonts/TTF/DejaVuSans.ttf'
                ]
                
                # Для Windows
                windows_fonts = [
                    'C:/Windows/Fonts/arial.ttf',
                    'C:/Windows/Fonts/calibri.ttf',
                    'C:/Windows/Fonts/tahoma.ttf'
                ]
                
                all_fonts = mac_fonts + linux_fonts + windows_fonts
                
                for font_path in all_fonts:
                    if os.path.exists(font_path):
                        try:
                            font_name = Path(font_path).stem.replace(' ', '')
                            # Обрабатываем .ttc файлы (коллекции шрифтов)
                            if font_path.endswith('.ttc'):
                                pdfmetrics.registerFont(TTFont(font_name, font_path, subfontIndex=0))
                            else:
                                pdfmetrics.registerFont(TTFont(font_name, font_path))
                            print(f"✅ Зарегистрирован шрифт: {font_name} из {font_path}")
                            break
                        except Exception as font_error:
                            print(f"⚠️ Не удалось зарегистрировать шрифт {font_path}: {font_error}")
                            continue
                
                # Если не удалось найти системный шрифт, используем стандартный
                if not font_name:
                    print("⚠️ Системные шрифты не найдены, используем стандартный Helvetica")
                    font_name = 'Helvetica'
                        
            except Exception as e:
                print(f"⚠️ Ошибка регистрации шрифта: {e}")
                font_name = 'Helvetica'
            
            doc = SimpleDocTemplate(str(file_path), pagesize=A4)
            styles = getSampleStyleSheet()
            
            # Создаем стили с поддержкой UTF-8
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName=font_name,
                fontSize=16,
                spaceAfter=30,
                alignment=TA_LEFT,
                encoding='utf-8'
            )
            
            info_style = ParagraphStyle(
                'InfoStyle',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=10,
                spaceAfter=20,
                encoding='utf-8'
            )
            
            time_style = ParagraphStyle(
                'TimeStyle',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=9,
                spaceAfter=5,
                encoding='utf-8'
            )
            
            speaker_style = ParagraphStyle(
                'SpeakerStyle',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=11,
                spaceAfter=5,
                encoding='utf-8'
            )
            
            text_style = ParagraphStyle(
                'TextStyle',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=11,
                spaceAfter=15,
                encoding='utf-8'
            )
            
            story = []
            
            # Заголовок
            title_text = f"Транскрипция: {filename}"
            story.append(Paragraph(title_text, title_style))
            
            # Информация о файле
            info_text = f"""
            <b>Файл:</b> {filename}<br/>
            <b>Дата создания:</b> {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}<br/>
            <b>ID задачи:</b> {task_id}
            """
            story.append(Paragraph(info_text, info_style))
            story.append(Spacer(1, 20))
            
            # Добавляем сегменты
            for segment in segments:
                # Временная метка
                time_text = f"[{format_time_srt(segment['start'])} - {format_time_srt(segment['end'])}]"
                story.append(Paragraph(time_text, time_style))
                
                # Спикер (если есть)
                if 'speaker' in segment and segment['speaker'] is not None:
                    speaker_text = f"<b>Спикер {segment['speaker']}:</b>"
                    story.append(Paragraph(speaker_text, speaker_style))
                
                # Текст - правильно обрабатываем HTML entities но сохраняем кириллицу
                text = segment['text']
                # Экранируем только HTML специальные символы, кириллицу оставляем как есть
                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                story.append(Paragraph(text, text_style))
            
            doc.build(story)
            return str(file_path)
            
        except Exception as e:
            print(f"❌ Ошибка создания PDF: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    @classmethod
    def generate_all_formats(cls, segments: List[Dict], task_id: str, filename: str, temp: bool = False) -> Dict[str, str]:
        """Генерация всех форматов субтитров"""
        subtitle_files = {}
        
        # Стандартные форматы
        formats = [
            ('srt', cls.generate_srt),
            ('vtt', cls.generate_vtt),
            ('tsv', cls.generate_tsv),
        ]
        
        # Дополнительные форматы (если доступны)
        if DOCX_AVAILABLE:
            formats.append(('docx', cls.generate_docx))
        
        if PDF_AVAILABLE:
            formats.append(('pdf', cls.generate_pdf))
        
        for format_name, generator_func in formats:
            try:
                file_path = generator_func(segments, task_id, filename, temp)
                if file_path:
                    subtitle_files[format_name] = file_path
                    print(f"✅ {format_name.upper()} файл создан: {Path(file_path).name}")
                else:
                    print(f"❌ {format_name.upper()} файл не создан")
            except Exception as e:
                print(f"❌ Ошибка создания {format_name.upper()} файла: {e}")
        
        return subtitle_files 